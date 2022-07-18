from docx import Document

doc = Document('微信号加好友时快速通过技巧.docx') #doc->docx
print(doc.paragraphs)

for p in doc.paragraphs:
    print(p.text)